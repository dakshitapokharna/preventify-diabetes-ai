"""
scripts/generate_compare_doc.py
Generate a formatted Word doc from the 3 compare-mode log files.
"""

import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).resolve().parent.parent
LOGS_DIR = BASE_DIR / "logs"
OUT_PATH = BASE_DIR / "Preventify_Compare_Output.docx"

LOG_FILES = [
    ("Knowledge-Based", LOGS_DIR / "model_compare_20260604_143158.json"),
    ("Lifestyle / Suggestion", LOGS_DIR / "model_compare_20260604_143433.json"),
    ("Medication-Related", LOGS_DIR / "model_compare_20260604_143855.json"),
]

BRAND_BLUE  = RGBColor(0x1A, 0x56, 0xDB)   # #1A56DB
BRAND_DARK  = RGBColor(0x11, 0x18, 0x27)   # #111827
LABEL_GREY  = RGBColor(0x6B, 0x72, 0x80)   # #6B7280
TABLE_HEADER_BG = "1A56DB"
TABLE_ROW_ALT   = "EFF6FF"


def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _para_spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def build_doc():
    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(title, before=0, after=60)
    run = title.add_run("Preventify — AI Compare Mode Output Report")
    run.bold      = True
    run.font.size = Pt(20)
    run.font.color.rgb = BRAND_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(sub, before=0, after=200)
    r = sub.add_run("Date: 04 June 2026     |     Pipeline: Phase 1 + RAG + Multi-Model Fan-out     |     Location default: Kerala, India")
    r.font.size = Pt(10)
    r.font.color.rgb = LABEL_GREY

    doc.add_paragraph()  # spacer

    # ── One section per question ──────────────────────────────────────────────
    for q_num, (category, log_path) in enumerate(LOG_FILES, start=1):
        with open(log_path, encoding="utf-8") as f:
            data = json.load(f)

        prompt  = data["prompt"].replace("\n  ▎ ", " ").strip().strip('"')
        results = data["results"]

        # Category chip
        cat_para = doc.add_paragraph()
        _para_spacing(cat_para, before=120, after=40)
        cat_run = cat_para.add_run(f"  Category: {category}  ")
        cat_run.bold           = True
        cat_run.font.size      = Pt(9)
        cat_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background via highlight hack — use shading on the paragraph
        pPr  = cat_para._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  TABLE_HEADER_BG)
        pPr.append(shd)

        # Question number + text
        q_heading = doc.add_paragraph()
        _para_spacing(q_heading, before=40, after=60)
        qn_run = q_heading.add_run(f"Q{q_num}.  ")
        qn_run.bold           = True
        qn_run.font.size      = Pt(13)
        qn_run.font.color.rgb = BRAND_BLUE
        qt_run = q_heading.add_run(prompt)
        qt_run.bold           = True
        qt_run.font.size      = Pt(13)
        qt_run.font.color.rgb = BRAND_DARK

        # ── Results table ─────────────────────────────────────────────────────
        col_widths = [Inches(1.0), Inches(1.6), Inches(3.8), Inches(0.6), Inches(0.7), Inches(0.7)]
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        headers   = ["Provider", "Model", "Response", "Latency (s)", "Input\nTokens", "Output\nTokens"]
        for i, (cell, hdr) in enumerate(zip(hdr_cells, headers)):
            cell.width = col_widths[i]
            _set_cell_bg(cell, TABLE_HEADER_BG)
            p   = cell.paragraphs[0]
            run = p.add_run(hdr)
            run.bold           = True
            run.font.size      = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        # Data rows
        for row_idx, res in enumerate(results):
            row_cells = table.add_row().cells
            bg = TABLE_ROW_ALT if row_idx % 2 == 0 else "FFFFFF"

            values = [
                res.get("provider", ""),
                res.get("model", ""),
                res.get("text", ""),
                str(res.get("latency_s", "")),
                str(res.get("input_tokens", "")),
                str(res.get("output_tokens", "")),
            ]
            for i, (cell, val) in enumerate(zip(row_cells, values)):
                cell.width = col_widths[i]
                _set_cell_bg(cell, bg)
                p   = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(9)
                if i == 2:   # response column — slightly larger
                    run.font.size = Pt(9)
                if i in (0, 1):
                    run.bold = True
                    run.font.color.rgb = BRAND_DARK

        doc.add_paragraph()  # spacer between questions

    # ── Footer note ───────────────────────────────────────────────────────────
    note = doc.add_paragraph()
    _para_spacing(note, before=120, after=0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "All responses generated by the Preventify server pipeline "
        "(Phase 1 context analysis → RAG retrieval → multi-model fan-out). "
        "Location defaulted to Kerala, India. For internal review only."
    )
    nr.font.size      = Pt(8)
    nr.font.color.rgb = LABEL_GREY
    nr.italic         = True

    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
